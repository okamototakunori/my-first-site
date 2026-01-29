
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# ==========================================
# 1. Generate Lesson Plan (DOCX)
# ==========================================

def create_lesson_plan():
    doc = docx.Document()

    # Page Setup (Margins 25mm)
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # Style Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Mincho'
    font.size = Pt(10.5)

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.space_after = Pt(6)
        p.space_before = Pt(12)

    def add_text(text):
        doc.add_paragraph(text)

    # --- Content ---

    # Title
    title = doc.add_paragraph('理科 学習指導案')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)

    # 1. 基本情報
    add_heading('1. 基本情報')
    add_text('日時： 令和8年2月3日（月） 第3校時')
    add_text('場所： 3年1組教室')
    add_text('学級： 第3学年1組（30名）')
    add_text('授業者： 〇〇 〇〇')

    # 2. 単元名
    add_heading('2. 単元名')
    add_text('音の性質')

    # 3. 単元の目標
    add_heading('3. 単元の目標')
    doc.add_paragraph('・【知識及び技能】音が出たり伝わったりするときには物が震えていることや、音の大きさが変わるときには物の震え方が変わることを理解している。', style='List Bullet')
    doc.add_paragraph('・【思考力・判断力・表現力等】音が出ているときや伝わっているときの物の様子について追究する中で、差異点や共通点を基に、音の性質についての問題を見いだし、表現している。', style='List Bullet')
    doc.add_paragraph('・【学びに向かう力・人間性等】音の性質について進んで調べようとし、学習したことを日常生活に活かそうとしている。', style='List Bullet')

    # 4. 単元の評価規準
    add_heading('4. 単元の評価規準')
    table_criteria = doc.add_table(rows=2, cols=3)
    table_criteria.style = 'Table Grid'
    hdr_cells = table_criteria.rows[0].cells
    hdr_cells[0].text = '知識・技能'
    hdr_cells[1].text = '思考・判断・表現'
    hdr_cells[2].text = '主体的に学習に取り組む態度'
    
    row_cells = table_criteria.rows[1].cells
    row_cells[0].text = '音が出たり伝わったりするときには物が震えていることや、音の大きさと震え方の関係を理解している。'
    row_cells[1].text = '音の性質について、差異点や共通点を基に問題を見いだし、実験結果を基に考察し表現している。'
    row_cells[2].text = '音の性質について進んで調べ、学習したことを日常生活と関連付けようとしている。'

    # 5. 単元設定の理由
    add_heading('5. 単元設定の理由')
    add_text('【児童観】')
    add_text('　本学級の児童は、理科の学習に対して意欲的であり、観察や実験を楽しみにしている。しかし、現象の背後にある科学的な原理を考察することには個人差がある。音については日常的に触れているものの、「なぜ音が聞こえるのか」「音はどのように伝わるのか」といった科学的な視点で捉えた経験は少ない。')
    add_text('【教材観】')
    add_text('　本単元は、学習指導要領の「A物質・エネルギー」の「(3)音の性質」に関わるものである。音が物の震えによって生じ、伝わることを体験的に理解することは、中学校での「音の伝わり方と音の性質」の学習につながる重要な基礎となる。太鼓、トライアングル、糸電話など、児童が実際に触れて確かめられる教材を用いることで、科学的な見方・考え方を育てることができる。')
    add_text('【指導観】')
    add_text('　指導にあたっては、まず音を出す活動を通して、物が震えていることに気づかせる。その上で、音の大きさと震え方の関係を実験を通して確かめさせる。また、糸電話を用いて音の伝わり方を体験的に理解させる。グループでの観察・実験を重視し、互いの気づきを共有することで、科学的な思考力を育てる。')

    # 6. 単元指導計画
    add_heading('6. 単元指導計画（全5時間）')
    table_plan = doc.add_table(rows=6, cols=4)
    table_plan.style = 'Table Grid'
    
    # Headers
    headers = ['次', '時', '学習活動', '評価規準']
    for i, h in enumerate(headers):
        table_plan.rows[0].cells[i].text = h
        
    # Content
    plan_data = [
        ['1', '1', '・身の回りの音を出すものを調べる。\n・音が出ているときの物の様子を観察する。（本時）', '思①'],
        ['1', '2', '・太鼓やトライアングルなど、音が出ているときの震えを確かめる。\n・音が出るときには物が震えていることをまとめる。', '知①'],
        ['2', '3', '・音の大きさと震え方の関係を調べる実験を計画する。\n・実験を行い、結果を記録する。', '思②'],
        ['2', '4', '・実験結果を基に、音の大きさと震え方の関係を考察する。\n・音の大きさが変わると震え方が変わることをまとめる。', '知②'],
        ['3', '5', '・糸電話を作り、音の伝わり方を調べる。\n・学習したことを振り返り、日常生活と関連付ける。', '態①'],
    ]
    
    for row_idx, row_data in enumerate(plan_data):
        cells = table_plan.rows[row_idx+1].cells
        for col_idx, text in enumerate(row_data):
            cells[col_idx].text = text

    # 7. 本時の指導
    add_heading('7. 本時の指導（第1時）')
    add_text('1. 目標： 音が出ているときの物の様子について、差異点や共通点を基に問題を見いだし、表現することができる。')
    add_text('2. 展開：')

    # 8. 本時の展開（詳細）
    table_detail = doc.add_table(rows=1, cols=6)
    table_detail.style = 'Table Grid'
    
    # Headers
    detail_headers = ['過程', '学習活動', '主発問・指示', '予想される児童の反応', '指導上の留意点・支援', '評価']
    for i, h in enumerate(detail_headers):
        table_detail.rows[0].cells[i].text = h
        
    # Data Rows
    detail_data = [
        [
            '導入\n(5分)',
            '1. 身の回りの音について話し合う\n・どんな音が聞こえるか\n・音はどうやって出るのか',
            '「今、どんな音が聞こえますか？」\n「音はどうやって出ると思いますか？」',
            '・時計の音が聞こえる\n・外から車の音が聞こえる\n・物をたたくと音が出る',
            '・児童の生活経験を引き出す\n・音に対する素朴な考えを把握する',
            ''
        ],
        [
            '展開1\n(15分)',
            '2. 音が出ているときの物の様子を観察する\n・太鼓をたたく\n・トライアングルをたたく\n・音叉を鳴らす\n・観察したことを記録する',
            '「太鼓をたたいたとき、太鼓の様子をよく観察してみましょう」\n「手で触ってみると、どんな感じがしますか？」',
            '・太鼓の皮が動いている\n・ビリビリする感じがする\n・トライアングルも震えている\n・音叉も震えている',
            '・グループに太鼓、トライアングル、音叉を配布\n・安全に配慮して観察させる\n・手で触れて確かめさせる\n・水を入れた容器に音叉を入れて水しぶきを観察させる',
            '思①'
        ],
        [
            '展開2\n(15分)',
            '3. 観察結果を発表し、共通点を見つける\n・各グループの観察結果を発表\n・共通点を話し合う\n・問題を見いだす',
            '「どの楽器でも、共通していることはありますか？」\n「音が出ているとき、物はどうなっていましたか？」',
            '・どれも震えていた\n・音が出ているときは震えている\n・大きな音のときは大きく震えていた',
            '・発表を板書で整理する\n・共通点に着目させる\n・「音の大きさと震え方の関係」という新たな問題を引き出す',
            '思①'
        ],
        [
            'まとめ\n(10分)',
            '4. 本時のまとめと振り返り\n・分かったことをまとめる\n・次時の学習内容を知る',
            '「今日の学習で分かったことをノートに書きましょう」\n「次の時間は、もっと詳しく調べていきます」',
            '・音が出ているときは物が震えている\n・音の大きさと震え方の関係を調べたい',
            '・キーワード「震え」を確認\n・次時への意欲を高める',
            ''
        ]
    ]

    for row_data in detail_data:
        row = table_detail.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # Set column widths
    for row in table_detail.rows:
        row.cells[0].width = Mm(20)
        row.cells[1].width = Mm(35)
        row.cells[2].width = Mm(35)
        row.cells[3].width = Mm(30)
        row.cells[4].width = Mm(30)
        row.cells[5].width = Mm(10)

    # 9. 板書計画
    add_heading('9. 板書計画')
    board_plan = """
 ____________________________________________________________________
|  2月3日  音の性質                                                  |
|                                                                    |
|  めあて： 音が出ているとき、物はどうなっているのだろうか          |
|                                                                    |
|  【観察したこと】                                                  |
|   太鼓        → 皮が震えていた、手で触るとビリビリした            |
|   トライアングル → 震えていた                                      |
|   音叉        → 震えていた、水に入れると水しぶきが出た            |
|                                                                    |
|  【共通点】                                                        |
|   ・どれも震えていた                                              |
|   ・音が出ているときは、物が震えている                            |
|                                                                    |
|  【新しい問題】                                                    |
|   音の大きさが変わると、震え方はどうなるのだろうか？              |
|                                                                    |
|  まとめ                                                            |
|   音が出ているとき、物は震えている。                              |
|____________________________________________________________________|
    """
    paragraph = doc.add_paragraph(board_plan)
    paragraph.style.font.name = 'Courier New'
    paragraph.style.font.size = Pt(9)

    doc.save('../指導案_理科_音の性質.docx')
    print("理科指導案（DOCX）を生成しました。")

# ==========================================
# 2. Generate Rubric (XLSX)
# ==========================================

def create_rubric():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "評価基準"

    # Define Data
    headers = ["評価の観点", "評価規準", "A（十分満足できる）", "B（おおむね満足できる）", "C（努力を要する）", "Cへの手立て"]
    data = [
        [
            "知識・技能①",
            "音が出ているときには物が震えていることを理解している。",
            "音が出ているときには物が震えていることを、複数の具体例を挙げて説明できる。",
            "音が出ているときには物が震えていることを理解し、説明できる。",
            "音が出ているときの物の様子を理解できていない。",
            "実際に楽器に触れさせ、震えを手で確かめさせる。"
        ],
        [
            "知識・技能②",
            "音の大きさが変わるときには物の震え方が変わることを理解している。",
            "音の大きさと震え方の関係を、実験結果を基に詳しく説明できる。",
            "音の大きさが変わると震え方が変わることを理解している。",
            "音の大きさと震え方の関係を理解できていない。",
            "実験を再度行い、震え方の違いを視覚的に確認させる。"
        ],
        [
            "思考・判断・表現①",
            "音の性質について、差異点や共通点を基に問題を見いだし、表現している。",
            "観察結果から共通点を見いだし、新たな問題を自ら設定して表現できる。",
            "観察結果から共通点を見いだし、問題を表現できる。",
            "観察結果から共通点を見いだすことが難しい。",
            "観察結果を整理する視点を示し、比較させる。"
        ],
        [
            "思考・判断・表現②",
            "音の大きさと震え方の関係について、実験結果を基に考察し表現している。",
            "実験結果を基に、音の大きさと震え方の関係を論理的に考察し、分かりやすく表現できる。",
            "実験結果を基に、音の大きさと震え方の関係を考察し表現できる。",
            "実験結果から考察することが難しい。",
            "実験結果を表にまとめさせ、規則性に気づかせる。"
        ],
        [
            "主体的に学習に取り組む態度①",
            "音の性質について進んで調べ、学習したことを日常生活と関連付けようとしている。",
            "音の性質について自ら進んで調べ、日常生活の様々な場面と関連付けて考えようとしている。",
            "音の性質について進んで調べ、日常生活と関連付けようとしている。",
            "学習への参加が消極的である。",
            "身近な音の例を示し、興味・関心を高める。"
        ]
    ]

    # Write Headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    # Write Data
    for row_num, row_data in enumerate(data, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    # Adjust Column Widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 25
    ws.column_dimensions['F'].width = 25

    wb.save('../評価基準_理科_音の性質.xlsx')
    print("評価基準（XLSX）を生成しました。")

if __name__ == "__main__":
    try:
        create_lesson_plan()
        create_rubric()
        print("すべてのファイルを作成しました。")
    except Exception as e:
        print(f"エラー: {e}")
