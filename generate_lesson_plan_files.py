
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# ==========================================
# 1. Generate Lesson Plan (DOCX)
# ==========================================

def create_lesson_plan():
    doc = docx.Document()

    # Page Setup (Margins 25mm)
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # Style Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Mincho' # Or Hiragino Mincho ProN on Mac if available, relying on Word's fallback or substitution
    font.size = Pt(10.5)

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.space_after = Pt(6)
        p.space_before = Pt(12)

    def add_text(text):
        doc.add_paragraph(text)

    # --- Content ---

    # Title
    title = doc.add_paragraph('国語科 学習指導案')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)

    # 1. 基本情報
    add_heading('1. 基本情報')
    add_text('日時： 令和8年1月27日（火） 第3校時')
    add_text('場所： 3年1組教室')
    add_text('学級： 第3学年1組（30名）')
    add_text('授業者： 〇〇 〇〇')

    # 2. 単元名
    add_heading('2. 単元名')
    add_text('文の組み立て（主語と述語）')

    # 3. 単元の目標
    add_heading('3. 単元の目標')
    doc.add_paragraph('・【知識及び技能】文には主語と述語があることや、主語と述語の関係、修飾と被修飾の関係について理解している。', style='List Bullet')
    doc.add_paragraph('・【思考力・判断力・表現力等】文の中における主語と述語との関係に注意して、文を正しく書いたり読んだりしている。', style='List Bullet')
    doc.add_paragraph('・【学びに向かう力・人間性等】進んで言葉の使い方を振り返り、学習したことを言語活動に活かそうとしている。', style='List Bullet')

    # 4. 単元の評価規準
    add_heading('4. 単元の評価規準')
    table_criteria = doc.add_table(rows=2, cols=3)
    table_criteria.style = 'Table Grid'
    hdr_cells = table_criteria.rows[0].cells
    hdr_cells[0].text = '知識・技能'
    hdr_cells[1].text = '思考・判断・表現'
    hdr_cells[2].text = '主体的に学習に取り組む態度'
    
    row_cells = table_criteria.rows[1].cells
    row_cells[0].text = '文の成分としての主語・述語の役割を理解し、見つけることができる。'
    row_cells[1].text = '主語と述語の対応関係を考え、適切に文を作ることができる。'
    row_cells[2].text = '言葉のきまりに関心を持ち、正しい文を書こうとしている。'

    # 5. 単元設定の理由
    add_heading('5. 単元設定の理由')
    add_text('【児童観】')
    add_text('　本学級の児童は、活発に発言し、学習に対して意欲的である。しかし、作文や日記などの記述を見ると、主語が抜けていたり、主語と述語がねじれていたりする文が散見される。文の構造を意識して書くことには個人差がある。')
    add_text('【教材観】')
    add_text('　本単元は、学習指導要領の「文の組立て」に関わるものである。文の基本構造である「何が（だれが）－どうする（どんなだ・なんだ）」という主語・述語の関係を理解することは、正確に情報を伝えるための基礎となる。')
    add_text('【指導観】')
    add_text('　指導にあたっては、短冊を用いた並べ替え活動や具体例を通して、視覚的に文の構造を捉えさせる。また、ペア学習を取り入れ、互いの作った文を読み合うことで、主語と述語の対応関係についての気づきを促す。')

    # 6. 単元指導計画
    add_heading('6. 単元指導計画（全2時間）')
    table_plan = doc.add_table(rows=3, cols=4)
    table_plan.style = 'Table Grid'
    
    # Headers
    headers = ['次', '時', '学習活動', '評価規準']
    for i, h in enumerate(headers):
        table_plan.rows[0].cells[i].text = h
        
    # Content
    plan_data = [
        ['1', '1', '・主語と述語の意味を知る。\n・簡単な文から主語と述語を見つける。（本時）', '知①'],
        ['1', '2', '・主語と述語のねじれに気づき、正しく直す。\n・主語と述語に気をつけて短文を作る。', '思①'],
    ]
    
    for row_idx, row_data in enumerate(plan_data):
        cells = table_plan.rows[row_idx+1].cells
        for col_idx, text in enumerate(row_data):
            cells[col_idx].text = text

    # 7. 本時の指導
    add_heading('7. 本時の指導')
    add_text('1. 目標： 文の成分である主語と述語の役割を理解し、文の中からそれらを見つけることができる。')
    add_text('2. 展開：')

    # 8. 本時の展開（詳細）
    table_detail = doc.add_table(rows=1, cols=6)
    table_detail.style = 'Table Grid'
    
    # Headers
    detail_headers = ['過程', '学習活動', '主発問・指示', '児童の反応', '指導・支援', '評価']
    for i, h in enumerate(detail_headers):
        table_detail.rows[0].cells[i].text = h
        
    # Data Rows
    detail_data = [
        [
            '導入\n(5分)',
            '1. 文のパズルをする\n・ばらばらの言葉を並べて文にする\n「犬が」「走る」',
            '「このカードを並べて文を作りましょう」\n「どちらが『だれが』で、どちらが『どうする』ですか？」',
            '・「犬が」が前だと思う\n・「犬が 走る」で文になる',
            '・簡単な文提示し、文の成立には「何が」と「どうする」が必要なことに気づかせる。',
            ''
        ],
        [
            '展開1\n(15分)',
            '2. 主語と述語を知る\n・「何が（だれが）」にあたる言葉＝主語\n・「どうする（どんなだ・なんだ）」にあたる言葉＝述語\n・教科書の例文で確認する',
            '「文の中の『何が（だれが）』にあたる部分を主語と言います」\n「『どうする（どんなだ・なんだ）』にあたる部分を述語と言います」',
            '・「〜は」も主語かな？\n・述語は文の最後にあることが多いな',
            '・主語には赤線、述語には青線を引くなどして視覚的に区別する。\n・「〜は、〜が」が主語の目印になることを助言する。',
            '知①'
        ],
        [
            '展開2\n(20分)',
            '3. 練習問題に取り組む\n・用意された文から主語と述語を見つける\n・ペアで確認し合う',
            '「次の文の主語に赤線、述語に青線を引きましょう」\n「隣の人と答え合わせをしましょう」',
            '・「赤い・花が・さいた」…どれが主語？\n・「花が」が主語だ。',
            '・修飾語（「赤い」など）と混同しないよう、「何が」「どうした」という問いかけに戻って考えさせる。\n・机間巡視で個別に支援する。',
            '知①'
        ],
        [
            'まとめ\n(5分)',
            '4. 本時のまとめと振り返り',
            '「今日の学習で分かったことを書きましょう」',
            '・主語と述語という名前を覚えた。\n・文には主語と述語があることが分かった。',
            '・次時は自分で文を作ることを予告する。',
            ''
        ]
    ]

    for row_data in detail_data:
        row = table_detail.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # Set column widths (approximate relative widths)
    for row in table_detail.rows:
        row.cells[0].width = Mm(20) # 過程
        row.cells[1].width = Mm(35) # 活動
        row.cells[2].width = Mm(35) # 発問
        row.cells[3].width = Mm(30) # 反応
        row.cells[4].width = Mm(30) # 支援
        row.cells[5].width = Mm(10) # 評価

    # 9. 板書計画
    add_heading('9. 板書計画')
    board_plan = """
 ____________________________________________________________________
|  1月27日  文の組み立て                                             |
|                                                                    |
|  めあて： 文の中の「主語」と「述語」を見つけよう                    |
|                                                                    |
|  [例文カード]                                                      |
|   犬が   走る                                                      |
|   ~~~~   ~~~~                                                      |
|  (だれが) (どうする)                                               |
|    ↓        ↓                                                      |
|   主語     述語                                                    |
|                                                                    |
|  ポイント                                                          |
|  ・主語 … 「だれが」「なにが」にあたる言葉  (赤線 ____________ ) |
|  ・述語 … 「どうする」「どんなだ」「なんだ」 (青線 ~~~~~~~~~~~~ ) |
|                                                                    |
|  [練習]                                                            |
|   1. 赤い  花が   さいた。                                         |
|            ^^^^   ~~~~~~                                           |
|            主語    述語                                            |
|                                                                    |
|  振り返り                                                          |
|  ・主語と述語を見つけるときは、「何が」「どうする」と問いかけると  |
|    わかりやすい。                                                  |
|____________________________________________________________________|
    """
    paragraph = doc.add_paragraph(board_plan)
    paragraph.style.font.name = 'Courier New' # Monospace for ASCII art board plan
    paragraph.style.font.size = Pt(9)

    doc.save('指導案_国語_文の組み立て.docx')
    print("DOCX generated.")

# ==========================================
# 2. Generate Rubric (XLSX)
# ==========================================

def create_rubric():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "評価基準"

    # Define Data
    headers = ["評価の観点", "評価規準", "A（十分満足できる）", "B（おおむね満足できる）", "C（努力を要する）", "Cへの手立て"]
    data = [
        [
            "知識・技能①",
            "文の成分としての主語・述語の役割を理解し、見つけることができる。",
            "複雑な文（修飾語を含む文など）においても、主語と述語を正確に見つけることができる。",
            "簡単な単文において、主語と述語を見つけることができる。",
            "主語と述語の意味が理解できず、文の中から見つけることが難しい。",
            "「何が」「どうする」といった問いかけを個別に行い、文節ごとに区切って考えさせる。"
        ],
        [
            "思考・判断・表現①",
            "主語と述語の対応関係を考え、適切な文を作ることができる。",
            "主語と述語の呼応だけでなく、修飾語を用いてより詳しい文を作ることができる。",
            "主語と述語のねじれのない、正しい文を作ることができる。",
            "主語が抜けたり、主語と述語がねじれたりした文になってしまう。",
            "短冊カードを用いて、主語と述語のペアを作る活動を通して感覚を掴ませる。"
        ],
        [
            "主体的に学習に取り組む態度①",
            "言葉のきまりに関心を持ち、正しい文を書こうとしている。",
            "日常生活の言葉使いや、他の教科の記述においても、主語と述語の関係を意識しようとしている。",
            "学習活動において、進んで主語と述語を見つけたり、文を作ろうとしたりしている。",
            "活動への参加が消極的である。",
            "できたところを認め、スモールステップで自信を持たせる。"
        ]
    ]

    # Write Headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    # Write Data
    for row_num, row_data in enumerate(data, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    # Adjust Column Widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 25
    ws.column_dimensions['F'].width = 25

    wb.save('評価基準_国語_文の組み立て.xlsx')
    print("XLSX generated.")

if __name__ == "__main__":
    try:
        create_lesson_plan()
        create_rubric()
        print("All files created successfully.")
    except Exception as e:
        print(f"Error: {e}")
